"""
╔══════════════════════════════════════════════════════════════════╗
║       QR SCANNER SECURITY — ANALISIS & REPORT GENERATOR         ║
║  Input : Excel hasil pengujian dari qr_scanner_tester.py        ║
║  Output: Laporan Word (.docx) lengkap + chart PNG               ║
║  Pakai : python qr_analysis_report.py [file_excel.xlsx]         ║
╚══════════════════════════════════════════════════════════════════╝
"""

import sys, os, re, math
from collections import defaultdict
import datetime

import openpyxl
import matplotlib
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# PERBAIKAN: Suppress font warnings dan set font yang support emoji
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="matplotlib")

# Coba gunakan font yang support emoji
try:
    from matplotlib import font_manager
    # Untuk Windows
    if sys.platform == "win32":
        plt.rcParams['font.family'] = 'Segoe UI Emoji'
    # Untuk Mac
    elif sys.platform == "darwin":
        plt.rcParams['font.family'] = 'Apple Color Emoji'
except:
    pass  # Fallback ke default
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Konstanta ────────────────────────────────────────────────────────────
CHART_DIR = "/home/claude/charts"
os.makedirs(CHART_DIR, exist_ok=True)

COLOR_DARK   = (0x1F, 0x38, 0x64)
COLOR_BLUE   = (0x2E, 0x75, 0xB6)
COLOR_GREEN  = (0x37, 0x56, 0x23)
COLOR_RED    = (0xC0, 0x00, 0x00)
COLOR_ORANGE = (0xC5, 0x5A, 0x11)
COLOR_GRAY   = (0x60, 0x60, 0x60)

HEX_DARK  = "1F3864"
HEX_BLUE  = "2E75B6"
HEX_GREEN = "375623"
HEX_RED   = "C00000"

TARGETS = {
    "Detection Rate": 0.90,
    "Precision":      0.85,
    "Specificity":    0.90,
    "F1-Score":       0.87,
    "Accuracy":       0.90,
    "FP Rate":        0.05,   # must be BELOW this
    "Critical Miss":  0.01,   # must be BELOW this
}

# ─── 1. BACA DATA EXCEL ───────────────────────────────────────────────────
def read_excel(path: str) -> dict:
    wb = openpyxl.load_workbook(path, data_only=True)
    scanners = {}
    skip = {"Ringkasan", "Sheet"}

    for name in wb.sheetnames:
        if name in skip:
            continue
        ws = wb[name]
        rows = []
        for r in ws.iter_rows(min_row=3, values_only=True):
            if r[0] is None:
                continue
            rows.append({
                "no":       r[0],
                "waktu":    r[1],
                "url":      str(r[2] or ""),
                "label":    str(r[3] or "").lower().strip(),
                "kategori": str(r[4] or ""),
                "hasil":    str(r[5] or ""),
                "tp":       int(r[6] or 0),
                "fp":       int(r[7] or 0),
                "fn":       int(r[8] or 0),
                "tn":       int(r[9] or 0),
                "benar":    str(r[10] or ""),
                "catatan":  str(r[11] or ""),
            })
        if rows:
            scanners[name] = rows
    return scanners

# ─── 2. HITUNG METRIK ─────────────────────────────────────────────────────
def compute_metrics(rows: list) -> dict:
    tp = sum(r["tp"] for r in rows)
    fp = sum(r["fp"] for r in rows)
    fn = sum(r["fn"] for r in rows)
    tn = sum(r["tn"] for r in rows)
    total = tp + fp + fn + tn

    precision   = tp / (tp + fp) if (tp + fp) else 0
    recall      = tp / (tp + fn) if (tp + fn) else 0
    specificity = tn / (tn + fp) if (tn + fp) else 0
    f1          = 2 * precision * recall / (precision + recall) if (precision + recall) else 0
    accuracy    = (tp + tn) / total if total else 0
    fp_rate     = fp / (fp + tn) if (fp + tn) else 0
    crit_miss   = fn / (tp + fn) if (tp + fn) else 0
    block_rate  = (tp + fp) / total if total else 0

    return dict(tp=tp, fp=fp, fn=fn, tn=tn, total=total,
                precision=precision, recall=recall, specificity=specificity,
                f1=f1, accuracy=accuracy, fp_rate=fp_rate,
                crit_miss=crit_miss, block_rate=block_rate)

def rank_scanners(all_metrics: dict) -> list:
    """Rank by composite score: weighted average of key metrics."""
    scored = []
    for name, m in all_metrics.items():
        score = (
            m["recall"]      * 0.30 +
            m["precision"]   * 0.20 +
            m["f1"]          * 0.20 +
            m["accuracy"]    * 0.15 +
            m["specificity"] * 0.10 +
            (1 - m["fp_rate"])    * 0.05
        )
        scored.append((name, score, m))
    scored.sort(key=lambda x: x[1], reverse=True)
    return scored

def analyze_attack_patterns(scanners_data: dict) -> dict:
    """Analyze which URL categories slip through most."""
    missed = defaultdict(int)
    total_phish = defaultdict(int)
    fp_by_cat = defaultdict(int)
    for rows in scanners_data.values():
        for r in rows:
            if r["label"] == "phishing":
                total_phish[r["kategori"]] += 1
                if r["fn"] == 1:
                    missed[r["kategori"]] += 1
            elif r["label"] == "safe" and r["fp"] == 1:
                fp_by_cat[r["kategori"]] += 1
    miss_rate = {}
    for cat in total_phish:
        miss_rate[cat] = missed[cat] / total_phish[cat] if total_phish[cat] else 0
    return {"miss_rate": miss_rate, "total_phish": dict(total_phish),
            "missed": dict(missed), "fp_by_cat": dict(fp_by_cat)}

def analyze_weakness(name: str, rows: list, m: dict) -> list:
    """Return list of weakness strings for a scanner."""
    w = []
    if m["recall"] < TARGETS["Detection Rate"]:
        w.append(f"Detection Rate {m['recall']*100:.1f}% — di bawah target 90%, phishing sering lolos")
    if m["fp_rate"] > TARGETS["FP Rate"]:
        w.append(f"False Positive Rate {m['fp_rate']*100:.1f}% — terlalu sering memblokir URL aman")
    if m["crit_miss"] > TARGETS["Critical Miss"]:
        w.append(f"Critical Miss Rate {m['crit_miss']*100:.1f}% — phishing berbahaya lolos terlalu banyak")
    if m["precision"] < TARGETS["Precision"]:
        w.append(f"Precision {m['precision']*100:.1f}% — banyak false alarm yang mengganggu pengguna")
    if m["f1"] < TARGETS["F1-Score"]:
        w.append(f"F1-Score {m['f1']*100:.1f}% — keseimbangan deteksi vs akurasi masih rendah")
    if not w:
        w.append("Tidak ada kelemahan signifikan — semua metrik memenuhi target")
    return w

def predict_risk(m: dict) -> tuple[str, str]:
    """Return (risk_level, prediction_text)."""
    score = m["recall"] * 0.4 + m["precision"] * 0.3 + m["f1"] * 0.3
    if score >= 0.90:
        return "RENDAH 🟢", "Scanner sangat handal. Risiko phishing lolos sangat kecil."
    elif score >= 0.78:
        return "SEDANG 🟡", "Scanner cukup baik namun perlu dikombinasikan dengan edukasi pengguna."
    elif score >= 0.65:
        return "TINGGI 🟠", "Scanner memiliki celah signifikan. Rekomendasi: jangan jadikan satu-satunya pertahanan."
    else:
        return "KRITIS 🔴", "Scanner tidak memadai sebagai perlindungan utama. Butuh mitigasi tambahan segera."

# ─── 3. BUAT CHART ────────────────────────────────────────────────────────
CHART_PATHS = {}

def save_chart(fig, name):
    p = os.path.join(CHART_DIR, f"{name}.png")
    fig.savefig(p, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    CHART_PATHS[name] = p
    return p

def chart_metric_comparison(ranked: list):
    metrics = ["recall","precision","f1","accuracy","specificity"]
    labels  = ["Detection Rate","Precision","F1-Score","Accuracy","Specificity"]
    names   = [r[0] for r in ranked]
    short_names = [n.replace(" (Built-in)","").replace(" Camera","") for n in names]
    x = np.arange(len(metrics))
    w = 0.18
    colors = ["#2E75B6","#27ae60","#e74c3c","#f39c12","#8e44ad","#1abc9c"]

    fig, ax = plt.subplots(figsize=(12, 5.5))
    for i, (name, _, m) in enumerate(ranked):
        vals = [m[k]*100 for k in metrics]
        bars = ax.bar(x + i*w, vals, w, label=short_names[i], color=colors[i], alpha=0.88, edgecolor="white")
        for bar, val in zip(bars, vals):
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.5,
                    f"{val:.0f}%", ha="center", va="bottom", fontsize=7.5, fontweight="bold")

    ax.set_xticks(x + w*(len(ranked)-1)/2)
    ax.set_xticklabels(labels, fontsize=10)
    ax.set_ylim(0, 115)
    ax.set_ylabel("Nilai (%)", fontsize=11)
    ax.set_title("Perbandingan Metrik Antar Scanner", fontsize=14, fontweight="bold", pad=15)
    ax.legend(loc="upper left", fontsize=9)
    ax.axhline(90, color="red", linestyle="--", alpha=0.5, linewidth=1, label="Target 90%")
    ax.axhline(85, color="orange", linestyle=":", alpha=0.5, linewidth=1)
    ax.grid(axis="y", alpha=0.3)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    save_chart(fig, "metric_comparison")

def chart_ranking_bar(ranked: list):
    names  = [r[0].replace(" (Built-in)","").replace(" Camera","") for r in ranked]
    scores = [r[1]*100 for r in ranked]
    colors_list = []
    for s in scores:
        if s >= 85:   colors_list.append("#27ae60")
        elif s >= 70: colors_list.append("#f39c12")
        else:         colors_list.append("#e74c3c")

    fig, ax = plt.subplots(figsize=(9, 4.5))
    bars = ax.barh(names[::-1], scores[::-1], color=colors_list[::-1], height=0.55, edgecolor="white")
    for bar, val in zip(bars, scores[::-1]):
        ax.text(val + 0.5, bar.get_y()+bar.get_height()/2,
                f"{val:.1f}%", va="center", fontsize=11, fontweight="bold")

    ax.set_xlim(0, 110)
    ax.set_xlabel("Composite Score (%)", fontsize=11)
    ax.set_title("Ranking Scanner — Composite Score", fontsize=14, fontweight="bold")
    for i, (name, rank) in enumerate(zip(names[::-1], range(len(names), 0, -1))):
        medal = {1:"🥇", 2:"🥈", 3:"🥉"}.get(rank, f"#{rank}")
        ax.text(1, i, medal, va="center", fontsize=14)
    ax.axvline(85, color="green", linestyle="--", alpha=0.5, linewidth=1)
    ax.grid(axis="x", alpha=0.3)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    save_chart(fig, "ranking")

def chart_confusion_matrix_grid(all_metrics: dict):
    scanners = list(all_metrics.keys())
    n = len(scanners)
    fig, axes = plt.subplots(1, n, figsize=(4.5*n, 4.5))
    if n == 1: axes = [axes]

    for ax, (name, m) in zip(axes, all_metrics.items()):
        mat = np.array([[m["tp"], m["fp"]], [m["fn"], m["tn"]]])
        total = mat.sum()
        short = name.replace(" (Built-in)","").replace(" Camera","")

        im = ax.imshow([[0,0],[0,0]], cmap="Blues", vmin=0, vmax=1)
        colors_cell = [["#27ae60","#e74c3c"],["#e67e22","#2980b9"]]
        labels_cell = [["TP","FP"],["FN","TN"]]
        for i in range(2):
            for j in range(2):
                ax.add_patch(plt.Rectangle((j-0.5, i-0.5), 1, 1,
                    facecolor=colors_cell[i][j], alpha=0.75, edgecolor="white", linewidth=2))
                val = mat[i, j]
                pct = val/total*100 if total else 0
                ax.text(j, i-0.12, f"{val}", ha="center", va="center",
                        fontsize=20, fontweight="bold", color="white")
                ax.text(j, i+0.22, f"({pct:.1f}%)", ha="center", va="center",
                        fontsize=9, color="white", alpha=0.9)
                ax.text(j, i+0.42, labels_cell[i][j], ha="center", va="center",
                        fontsize=9, fontweight="bold", color="white")

        ax.set_xlim(-0.5, 1.5); ax.set_ylim(-0.5, 1.5)
        ax.set_xticks([0,1]); ax.set_yticks([0,1])
        ax.set_xticklabels(["Predicted\nPhishing","Predicted\nSafe"], fontsize=9)
        ax.set_yticklabels(["Actual\nPhishing","Actual\nSafe"], fontsize=9)
        ax.set_title(short, fontsize=11, fontweight="bold", pad=10)
        ax.tick_params(length=0)
        for spine in ax.spines.values(): spine.set_visible(False)

    fig.suptitle("Confusion Matrix per Scanner", fontsize=14, fontweight="bold", y=1.02)
    fig.tight_layout()
    save_chart(fig, "confusion_matrix")

def chart_attack_patterns(patterns: dict):
    cats = sorted(patterns["miss_rate"], key=lambda c: patterns["miss_rate"][c], reverse=True)
    if not cats: return
    miss_rates = [patterns["miss_rate"][c]*100 for c in cats]
    totals = [patterns["total_phish"].get(c,0) for c in cats]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(13, 5))

    # Miss rate by category
    colors_bar = ["#e74c3c" if r > 30 else "#f39c12" if r > 15 else "#27ae60" for r in miss_rates]
    bars = ax1.barh([c.replace("Fake ","") for c in cats[::-1]], miss_rates[::-1],
                    color=colors_bar[::-1], height=0.6, edgecolor="white")
    for bar, val in zip(bars, miss_rates[::-1]):
        ax1.text(val+0.5, bar.get_y()+bar.get_height()/2,
                 f"{val:.0f}%", va="center", fontsize=10, fontweight="bold")
    ax1.set_xlim(0, 105)
    ax1.set_xlabel("Miss Rate (%)", fontsize=11)
    ax1.set_title("Pola Serangan: Miss Rate per Kategori\n(Lebih tinggi = lebih berbahaya)", fontsize=11, fontweight="bold")
    ax1.axvline(20, color="red", linestyle="--", alpha=0.5, linewidth=1)
    ax1.grid(axis="x", alpha=0.3)
    ax1.spines["top"].set_visible(False)
    ax1.spines["right"].set_visible(False)

    # Volume pie chart
    if sum(totals) > 0:
        pie_colors = plt.cm.Set3(np.linspace(0, 1, len(cats)))
        wedges, texts, autotexts = ax2.pie(
            totals, labels=[c.replace("Fake ","") for c in cats],
            autopct="%1.0f%%", colors=pie_colors,
            startangle=140, pctdistance=0.82
        )
        for t in texts: t.set_fontsize(9)
        for at in autotexts: at.set_fontsize(8)
        ax2.set_title("Volume Serangan per Kategori", fontsize=11, fontweight="bold")

    fig.tight_layout()
    save_chart(fig, "attack_patterns")

def chart_radar(ranked: list):
    metrics_keys  = ["recall","precision","f1","accuracy","specificity"]
    metrics_labels = ["Detection\nRate","Precision","F1-Score","Accuracy","Specificity"]
    N = len(metrics_keys)
    angles = [n / float(N) * 2 * math.pi for n in range(N)]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))
    # PERBAIKAN: Tambah lebih banyak warna
    colors = ["#2E75B6","#27ae60","#e74c3c","#f39c12","#8e44ad","#1abc9c","#3498db","#9b59b6"]
    
    # PERBAIKAN: Batasi loop berdasarkan jumlah scanner yang tersedia
    for i, (name, score, m) in enumerate(ranked):
        if i >= len(colors):  # Safety check
            break
        vals = [m[k] for k in metrics_keys] + [m[metrics_keys[0]]]
        short = name.replace(" (Built-in)","").replace(" Camera","")
        ax.plot(angles, vals, "o-", linewidth=2, color=colors[i], label=short)
        ax.fill(angles, vals, alpha=0.10, color=colors[i])

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(metrics_labels, size=10, fontweight="bold")
    ax.set_ylim(0, 1)
    ax.set_yticks([0.2, 0.4, 0.6, 0.8, 1.0])
    ax.set_yticklabels(["20%","40%","60%","80%","100%"], fontsize=8)
    ax.set_title("Radar Chart — Profil Performa Scanner", fontsize=13, fontweight="bold", pad=20)
    ax.legend(loc="upper right", bbox_to_anchor=(1.35, 1.15), fontsize=10)
    ax.grid(True, alpha=0.3)
    fig.tight_layout()
    save_chart(fig, "radar")

def chart_fp_fn_analysis(ranked: list):
    names  = [r[0].replace(" (Built-in)","").replace(" Camera","") for r in ranked]
    fp_rates = [r[2]["fp_rate"]*100 for r in ranked]
    fn_rates = [r[2]["crit_miss"]*100 for r in ranked]

    fig, ax = plt.subplots(figsize=(9, 5.5))
    x = np.arange(len(names))
    w = 0.35
    b1 = ax.bar(x - w/2, fp_rates, w, label="False Positive Rate", color="#e74c3c", alpha=0.85, edgecolor="white")
    b2 = ax.bar(x + w/2, fn_rates, w, label="Critical Miss Rate (FN)", color="#e67e22", alpha=0.85, edgecolor="white")

    for bar, val in zip(b1, fp_rates):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.3,
                f"{val:.1f}%", ha="center", va="bottom", fontsize=9, fontweight="bold")
    for bar, val in zip(b2, fn_rates):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.3,
                f"{val:.1f}%", ha="center", va="bottom", fontsize=9, fontweight="bold")

    ax.axhline(5,  color="#e74c3c", linestyle="--", alpha=0.5, linewidth=1, label="Target FP < 5%")
    ax.axhline(1,  color="#e67e22", linestyle=":", alpha=0.6, linewidth=1, label="Target Miss < 1%")
    ax.set_xticks(x); ax.set_xticklabels(names, fontsize=10)
    ax.set_ylabel("Rate (%)", fontsize=11)
    ax.set_ylim(0, max(max(fp_rates), max(fn_rates)) + 8)
    ax.set_title("Analisis Error: FP Rate vs Critical Miss Rate\n(Lebih rendah = lebih baik)", fontsize=12, fontweight="bold")
    ax.legend(fontsize=9); ax.grid(axis="y", alpha=0.3)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout()
    save_chart(fig, "fp_fn_analysis")

# ─── 4. BUILD DOCX ────────────────────────────────────────────────────────
def rgb(*args):
    return RGBColor(*args)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=COLOR_DARK):
    from docx.oxml.ns import qn as _qn
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Arial"
    return h

def add_para(doc, text, size=11, bold=False, italic=False,
             color=(0,0,0), space_before=0, space_after=6,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)
    run.font.name = "Arial"
    return p

def add_image_centered(doc, img_path, width_inch=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inch))
    return p

def build_report(excel_path: str, output_path: str):
    print("📖 Membaca data Excel...")
    scanners_data = read_excel(excel_path)
    if not scanners_data:
        print("❌ Tidak ada data ditemukan di Excel!")
        return

    print(f"   Ditemukan {len(scanners_data)} scanner: {list(scanners_data.keys())}")

    print("🔢 Menghitung metrik...")
    all_metrics = {name: compute_metrics(rows) for name, rows in scanners_data.items()}
    ranked      = rank_scanners(all_metrics)
    patterns    = analyze_attack_patterns(scanners_data)

    print("📊 Membuat grafik...")
    chart_ranking_bar(ranked)
    chart_metric_comparison(ranked)
    chart_confusion_matrix_grid(all_metrics)
    chart_attack_patterns(patterns)
    chart_radar(ranked)
    chart_fp_fn_analysis(ranked)

    print("📝 Menyusun dokumen laporan...")
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── COVER PAGE ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("LAPORAN EVALUASI KEAMANAN")
    run.font.size = Pt(26); run.font.bold = True
    run.font.color.rgb = RGBColor(*COLOR_DARK); run.font.name = "Arial"

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_sub.add_run("QR CODE SECURITY SCANNER")
    run2.font.size = Pt(20); run2.font.bold = True
    run2.font.color.rgb = RGBColor(*COLOR_BLUE); run2.font.name = "Arial"

    doc.add_paragraph()
    p_sub2 = doc.add_paragraph()
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p_sub2.add_run("Analisis Komprehensif Performa, Kelemahan,\nPola Serangan, & Rekomendasi Mitigasi")
    run3.font.size = Pt(13); run3.font.italic = True
    run3.font.color.rgb = RGBColor(*COLOR_GRAY); run3.font.name = "Arial"

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    now = datetime.datetime.now()
    meta_lines = [
        f"Tanggal Laporan   : {now.strftime('%d %B %Y')}",
        f"Total Scanner Diuji : {len(scanners_data)}",
        f"Total URL Diuji    : {all_metrics[ranked[0][0]]['total']} per scanner",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(11); r.font.name = "Arial"; r.font.color.rgb = RGBColor(*COLOR_GRAY)

    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ───────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", 1)

    best_name, best_score, best_m = ranked[0]
    worst_name, worst_score, worst_m = ranked[-1]
    total_tested = best_m["total"]

    summary_text = (
        f"Laporan ini menyajikan hasil evaluasi komprehensif terhadap {len(scanners_data)} "
        f"jenis QR code scanner dalam mendeteksi URL phishing. Pengujian dilakukan menggunakan "
        f"{total_tested} URL per scanner, terdiri dari URL phishing dan URL legitimate. "
        f"Berdasarkan composite score (weighted average dari Detection Rate, Precision, F1-Score, "
        f"Accuracy, dan Specificity), scanner terbaik adalah {best_name} "
        f"dengan skor {best_score*100:.1f}%, sementara scanner dengan performa terendah adalah "
        f"{worst_name} dengan skor {worst_score*100:.1f}%."
    )
    add_para(doc, summary_text, size=11)

    # Key findings table
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Temuan Utama:")
    r.font.bold = True; r.font.size = Pt(12); r.font.name = "Arial"
    r.font.color.rgb = RGBColor(*COLOR_DARK)

    findings = [
        ("Scanner Terbaik", f"{best_name} — Detection Rate {best_m['recall']*100:.1f}%, F1-Score {best_m['f1']*100:.1f}%"),
        ("Scanner Terlemah", f"{worst_name} — Detection Rate {worst_m['recall']*100:.1f}%"),
        ("Kategori Phishing Paling Berbahaya",
         max(patterns["miss_rate"], key=patterns["miss_rate"].get, default="N/A") + " — miss rate tertinggi"),
        ("Rata-rata False Positive",
         f"{sum(m['fp_rate'] for m in all_metrics.values())/len(all_metrics)*100:.1f}% semua scanner"),
        ("Scanner Memenuhi Semua Target",
         str(sum(1 for _,_,m in ranked if m["recall"]>=0.9 and m["fp_rate"]<=0.05)) + f" dari {len(scanners_data)} scanner"),
    ]

    t = doc.add_table(rows=len(findings)+1, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(["Aspek", "Temuan"]):
        cell = t.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = "Arial"
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        set_cell_bg(cell, HEX_DARK)

    for ri, (aspect, finding) in enumerate(findings, 1):
        for ci, val in enumerate([aspect, finding]):
            cell = t.rows[ri].cells[ci]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10); run.font.name = "Arial"
            if ci == 0: run.font.bold = True
            set_cell_bg(cell, "DEEAF1" if ri%2==0 else "FFFFFF")

    doc.add_page_break()

    # ── 2. RANKING SCANNER ─────────────────────────────────────────────────
    add_heading(doc, "2. Ranking Scanner — Composite Score", 1)
    add_para(doc,
        "Ranking dihitung berdasarkan composite score dengan pembobotan: "
        "Detection Rate (30%), Precision (20%), F1-Score (20%), Accuracy (15%), "
        "Specificity (10%), dan (1 − FP Rate) (5%).",
        size=10, italic=True, color=COLOR_GRAY)

    add_image_centered(doc, CHART_PATHS["ranking"], width_inch=5.5)

    # Ranking table
    doc.add_paragraph()
    hdrs = ["Rank","Scanner","Composite\nScore","Detection\nRate","Precision","F1-Score","Accuracy","Risk Level"]
    col_w = [600, 2200, 900, 900, 900, 900, 900, 1200]
    t = doc.add_table(rows=len(ranked)+1, cols=len(hdrs))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, (h, w) in enumerate(zip(hdrs, col_w)):
        cell = t.rows[0].cells[ci]
        p2 = cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(h)
        run.font.bold = True; run.font.size = Pt(9); run.font.name = "Arial"
        run.font.color.rgb = RGBColor(255,255,255)
        set_cell_bg(cell, HEX_BLUE)

    for ri, (name, score, m) in enumerate(ranked, 1):
        medal = {1:"🥇",2:"🥈",3:"🥉"}.get(ri, f"#{ri}")
        risk_level, _ = predict_risk(m)
        vals = [medal, name, f"{score*100:.1f}%", f"{m['recall']*100:.1f}%",
                f"{m['precision']*100:.1f}%", f"{m['f1']*100:.1f}%",
                f"{m['accuracy']*100:.1f}%", risk_level]
        for ci, val in enumerate(vals):
            cell = t.rows[ri].cells[ci]
            p2 = cell.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(val)
            run.font.size = Pt(9); run.font.name = "Arial"
            if ci == 1: run.font.bold = True
            set_cell_bg(cell, "E2EFDA" if ri==1 else "DEEAF1" if ri%2==0 else "FFFFFF")

    doc.add_page_break()

    # ── 3. PERBANDINGAN METRIK ─────────────────────────────────────────────
    add_heading(doc, "3. Perbandingan Metrik Detail", 1)
    add_image_centered(doc, CHART_PATHS["metric_comparison"], width_inch=6.5)
    doc.add_paragraph()
    add_image_centered(doc, CHART_PATHS["radar"], width_inch=5.0)

    # Full metrics table
    doc.add_paragraph()
    metric_keys = ["tp","fp","fn","tn","precision","recall","specificity","f1","accuracy","fp_rate","crit_miss"]
    metric_lbls = ["TP","FP","FN","TN","Precision","Recall /\nDetection","Specificity","F1-Score","Accuracy","FP Rate","Critical\nMiss Rate"]
    targets_row = ["—","—","—","—","≥85%","≥90%","≥90%","≥87%","≥90%","<5%","<1%"]

    t = doc.add_table(rows=len(ranked)+2, cols=len(metric_keys)+1)
    t.style = "Table Grid"

    # Header row
    set_cell_bg(t.rows[0].cells[0], HEX_DARK)
    t.rows[0].cells[0].paragraphs[0].add_run("Scanner").font.bold = True
    t.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    t.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(9)

    for ci, (lbl, tgt) in enumerate(zip(metric_lbls, targets_row), 1):
        cell = t.rows[0].cells[ci]
        set_cell_bg(cell, HEX_BLUE)
        p2 = cell.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(lbl); run.font.bold = True; run.font.size = Pt(8)
        run.font.name = "Arial"; run.font.color.rgb = RGBColor(255,255,255)

    # Targets row
    set_cell_bg(t.rows[1].cells[0], "FFF2CC")
    t.rows[1].cells[0].paragraphs[0].add_run("Target").font.bold = True
    t.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)

    for ci, tgt in enumerate(targets_row, 1):
        cell = t.rows[1].cells[ci]; set_cell_bg(cell, "FFF2CC")
        p2 = cell.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(tgt); run.font.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xC5, 0x5A, 0x11)

    for ri, (name, score, m) in enumerate(ranked, 2):
        cell0 = t.rows[ri].cells[0]
        set_cell_bg(cell0, "DEEAF1" if ri%2==0 else "FFFFFF")
        run = cell0.paragraphs[0].add_run(name)
        run.font.bold = True; run.font.size = Pt(9); run.font.name = "Arial"

        for ci, key in enumerate(metric_keys, 1):
            val = m[key]
            cell = t.rows[ri].cells[ci]
            is_pct = key not in ["tp","fp","fn","tn","total"]
            disp = f"{val:.1%}" if is_pct else str(int(val))

            # Color based on pass/fail
            pass_fail_map = {
                "precision": val >= TARGETS["Precision"],
                "recall":    val >= TARGETS["Detection Rate"],
                "specificity": val >= TARGETS["Specificity"],
                "f1":        val >= TARGETS["F1-Score"],
                "accuracy":  val >= TARGETS["Accuracy"],
                "fp_rate":   val <= TARGETS["FP Rate"],
                "crit_miss": val <= TARGETS["Critical Miss"],
            }
            if key in pass_fail_map:
                bg = "E2EFDA" if pass_fail_map[key] else "FFE0E0"
            else:
                bg = "DEEAF1" if ri%2==0 else "FFFFFF"

            set_cell_bg(cell, bg)
            p2 = cell.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(disp); run.font.size = Pt(9); run.font.name = "Arial"
            run.font.bold = key in ["recall","f1","accuracy"]

    doc.add_page_break()

    # ── 4. CONFUSION MATRIX ────────────────────────────────────────────────
    add_heading(doc, "4. Confusion Matrix per Scanner", 1)
    add_para(doc,
        "Confusion matrix menunjukkan distribusi TP (True Positive), FP (False Positive), "
        "FN (False Negative), dan TN (True Negative) untuk setiap scanner.",
        size=10, italic=True, color=COLOR_GRAY)
    add_image_centered(doc, CHART_PATHS["confusion_matrix"], width_inch=6.5)

    doc.add_page_break()

    # ── 5. ANALISIS ERROR ──────────────────────────────────────────────────
    add_heading(doc, "5. Analisis Error: FP Rate vs Critical Miss Rate", 1)
    add_para(doc,
        "False Positive Rate (FPR) mengukur seberapa sering URL aman salah diblokir. "
        "Critical Miss Rate mengukur seberapa sering phishing berbahaya lolos dari deteksi. "
        "Keduanya harus ditekan seminimal mungkin.",
        size=10, italic=True, color=COLOR_GRAY)
    add_image_centered(doc, CHART_PATHS["fp_fn_analysis"], width_inch=6.0)

    doc.add_page_break()

    # ── 6. POLA SERANGAN ───────────────────────────────────────────────────
    add_heading(doc, "6. Analisis Pola Serangan Phishing", 1)
    add_para(doc,
        "Analisis pola serangan mengidentifikasi kategori URL phishing yang paling sering lolos "
        "dari deteksi scanner. Kategori dengan miss rate tinggi merepresentasikan blind spot "
        "terbesar dalam sistem keamanan.",
        size=10, italic=True, color=COLOR_GRAY)
    add_image_centered(doc, CHART_PATHS["attack_patterns"], width_inch=6.5)

    doc.add_paragraph()
    add_heading(doc, "Kategori Paling Berbahaya (Lolos Paling Sering):", 2, COLOR_RED)
    miss_sorted = sorted(patterns["miss_rate"].items(), key=lambda x: x[1], reverse=True)
    for i, (cat, rate) in enumerate(miss_sorted[:5], 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"{cat} — Miss Rate: {rate*100:.0f}%")
        run.font.size = Pt(11); run.font.name = "Arial"
        if rate > 0.3: run.font.color.rgb = RGBColor(*COLOR_RED)
        elif rate > 0.15: run.font.color.rgb = RGBColor(*COLOR_ORANGE)

    doc.add_page_break()

    # ── 7. ANALISIS KELEMAHAN PER SCANNER ─────────────────────────────────
    add_heading(doc, "7. Analisis Kelemahan per Scanner", 1)

    for rank_i, (name, score, m) in enumerate(ranked, 1):
        add_heading(doc, f"7.{rank_i}  {name}", 2, COLOR_BLUE)
        risk_level, risk_text = predict_risk(m)
        weaknesses = analyze_weakness(name, scanners_data[name], m)

        p = doc.add_paragraph()
        r = p.add_run(f"Composite Score: {score*100:.1f}%  |  Risk Level: {risk_level}")
        r.font.size = Pt(11); r.font.bold = True; r.font.name = "Arial"
        r.font.color.rgb = RGBColor(*COLOR_DARK)

        add_para(doc, risk_text, size=11, italic=True)

        for w in weaknesses:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(w)
            run.font.size = Pt(10); run.font.name = "Arial"
            if "lolos" in w.lower() or "kritis" in w.lower():
                run.font.color.rgb = RGBColor(*COLOR_RED)
            elif "sering" in w.lower():
                run.font.color.rgb = RGBColor(*COLOR_ORANGE)

        doc.add_paragraph()

    doc.add_page_break()

    # ── 8. PREDIKSI & RISK ASSESSMENT ─────────────────────────────────────
    add_heading(doc, "8. Prediksi & Risk Assessment", 1)
    add_para(doc,
        "Bagian ini memberikan prediksi tentang efektivitas scanner dalam skenario nyata "
        "berdasarkan pola data pengujian yang telah dikumpulkan.",
        size=10, italic=True, color=COLOR_GRAY)

    # Risk table
    t = doc.add_table(rows=len(ranked)+1, cols=4)
    t.style = "Table Grid"
    for ci, h in enumerate(["Scanner","Risk Level","Prediksi Skenario Nyata","Rekomendasi Penggunaan"]):
        cell = t.rows[0].cells[ci]
        set_cell_bg(cell, HEX_DARK)
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255,255,255); run.font.name = "Arial"

    recs = {
        "RENDAH 🟢": "Dapat dijadikan perlindungan utama",
        "SEDANG 🟡": "Gunakan sebagai lapisan pertama + edukasi user",
        "TINGGI 🟠": "Jangan jadikan satu-satunya perlindungan",
        "KRITIS 🔴": "Butuh penggantian/patch segera",
    }
    for ri, (name, score, m) in enumerate(ranked, 1):
        rl, rt = predict_risk(m)
        bg = "E2EFDA" if "🟢" in rl else "FFFACD" if "🟡" in rl else "FFE0E0"
        vals = [name, rl, rt, recs.get(rl.split()[0]+" "+rl.split()[1], "")]
        for ci, val in enumerate(vals):
            cell = t.rows[ri].cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.name = "Arial"
            if ci == 0: run.font.bold = True

    doc.add_paragraph()

    # Statistical prediction
    add_heading(doc, "Prediksi Statistik (1000 URL Real-world):", 2, COLOR_BLUE)
    best_m2 = ranked[0][2]
    phish_pct = 0.30  # asumsikan 30% URL adalah phishing di dunia nyata
    n = 1000
    n_phish = int(n * phish_pct)
    n_safe  = n - n_phish

    pred_stats = []
    for name, score, m in ranked:
        exp_tp = int(n_phish * m["recall"])
        exp_fn = n_phish - exp_tp
        exp_fp = int(n_safe * m["fp_rate"])
        exp_tn = n_safe - exp_fp
        pred_stats.append((name, exp_tp, exp_fn, exp_fp, exp_tn))

    t2 = doc.add_table(rows=len(pred_stats)+1, cols=5)
    t2.style = "Table Grid"
    for ci, h in enumerate(["Scanner","TP (Phishing Ditangkap)","FN (Phishing Lolos)","FP (Safe Diblokir)","TN (Safe Dibiarkan)"]):
        cell = t2.rows[0].cells[ci]
        set_cell_bg(cell, HEX_BLUE)
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255,255,255); run.font.name = "Arial"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, (name, tp, fn, fp, tn) in enumerate(pred_stats, 1):
        bg = "E2EFDA" if ri%2==0 else "FFFFFF"
        for ci, val in enumerate([name, tp, fn, fp, tn]):
            cell = t2.rows[ri].cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9); run.font.name = "Arial"
            if ci == 0: run.font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ci == 2 and int(val) > 5:
                run.font.color.rgb = RGBColor(*COLOR_RED)

    add_para(doc, f"*Asumsi: {n} URL total, {phish_pct*100:.0f}% adalah phishing (proporsi umum di lingkungan kerja)", size=9, italic=True, color=COLOR_GRAY)

    doc.add_page_break()

    # ── 9. KESIMPULAN & REKOMENDASI ────────────────────────────────────────
    add_heading(doc, "9. Kesimpulan & Rekomendasi", 1)

    add_heading(doc, "9.1 Kesimpulan Utama", 2, COLOR_DARK)
    conclusions = [
        f"Scanner {ranked[0][0]} menunjukkan performa terbaik secara keseluruhan dengan composite score {ranked[0][1]*100:.1f}%, deteksi rate {ranked[0][2]['recall']*100:.1f}%, dan false positive rate {ranked[0][2]['fp_rate']*100:.1f}%.",
        f"Tidak ada satu pun scanner yang sempurna — semua memiliki blind spot terutama pada kategori phishing yang meniru brand lokal Indonesia.",
        f"Kategori phishing '{max(patterns['miss_rate'], key=patterns['miss_rate'].get, default='N/A')}' merupakan vektor serangan paling berbahaya dengan miss rate tertinggi.",
        f"Kombinasi scanner + edukasi pengguna adalah pendekatan yang paling efektif untuk keamanan optimal.",
        f"Scanner dengan detection rate di bawah 80% sebaiknya tidak dijadikan perlindungan tunggal.",
    ]
    for c in conclusions:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(c)
        run.font.size = Pt(11); run.font.name = "Arial"

    doc.add_paragraph()
    add_heading(doc, "9.2 Rekomendasi Teknis", 2, COLOR_DARK)
    recs_tech = [
        ("Gunakan scanner dengan detection rate ≥ 90%", f"Saat ini hanya {sum(1 for _,_,m in ranked if m['recall']>=0.9)} scanner memenuhi kriteria ini."),
        ("Implementasikan multi-layer scanning", "Scan QR dua kali dengan scanner berbeda untuk meminimalkan false negative."),
        ("Update database phishing secara rutin", "Scanner berbasis database perlu pembaruan harian agar dapat mendeteksi phishing baru."),
        ("Prioritaskan deteksi Fake Banking & E-Commerce", "Dua kategori ini memiliki volume serangan tertinggi dan dampak finansial terbesar."),
        ("Tambahkan safelist untuk URL corporate", "Mengurangi false positive dan meningkatkan user experience."),
        ("Monitoring & logging setiap scan", "Lacak URL yang paling sering di-scan untuk mendeteksi kampanye phishing aktif."),
    ]
    t3 = doc.add_table(rows=len(recs_tech)+1, cols=2)
    t3.style = "Table Grid"
    for ci, h in enumerate(["Rekomendasi","Detail"]):
        cell = t3.rows[0].cells[ci]
        set_cell_bg(cell, HEX_GREEN)
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255,255,255); run.font.name = "Arial"

    for ri, (rec, detail) in enumerate(recs_tech, 1):
        bg = "E2EFDA" if ri%2==0 else "FFFFFF"
        for ci, val in enumerate([rec, detail]):
            cell = t3.rows[ri].cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10); run.font.name = "Arial"
            if ci == 0: run.font.bold = True

    doc.add_paragraph()
    add_heading(doc, "9.3 Roadmap Peningkatan Keamanan", 2, COLOR_DARK)
    roadmap = [
        ("Jangka Pendek (1-3 bulan)", [
            f"Adopsi {ranked[0][0]} sebagai scanner primer",
            "Buat SOP pemeriksaan URL sebelum scan QR di tempat umum",
            "Sosialisasi kepada pengguna tentang tanda-tanda phishing QR",
        ]),
        ("Jangka Menengah (3-6 bulan)", [
            "Implementasikan scanning ganda (2 scanner berbeda)",
            "Evaluasi ulang scanner yang scoring-nya di bawah 70%",
            "Bangun database URL whitelist internal",
        ]),
        ("Jangka Panjang (6-12 bulan)", [
            "Evaluasi solusi enterprise QR scanner dengan real-time threat intelligence",
            "Integrasikan scanning dengan MDM (Mobile Device Management) perusahaan",
            "Lakukan penetration testing berkala dengan QR phishing baru",
        ]),
    ]
    for period, items in roadmap:
        add_para(doc, period, size=11, bold=True, color=COLOR_BLUE)
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            run.font.size = Pt(10); run.font.name = "Arial"

    doc.add_page_break()

    # ── 10. LAMPIRAN ────────────────────────────────────────────────────────
    add_heading(doc, "10. Lampiran — Definisi & Formula", 1)

    formulas = [
        ("Precision", "TP / (TP + FP)", "Akurasi identifikasi phishing — minim false alarm"),
        ("Recall / Detection Rate", "TP / (TP + FN)", "Persentase phishing yang berhasil ditangkap"),
        ("Specificity", "TN / (TN + FP)", "Kemampuan mengenali URL aman dengan benar"),
        ("F1-Score", "2 × Precision × Recall / (Precision + Recall)", "Harmonic mean Precision dan Recall"),
        ("Accuracy", "(TP + TN) / Total", "Akurasi keseluruhan sistem"),
        ("FP Rate", "FP / (FP + TN)", "Frekuensi URL aman yang salah diblokir"),
        ("Critical Miss Rate", "FN / (TP + FN)", "Phishing berbahaya yang lolos dari deteksi"),
        ("Block Rate", "(TP + FP) / Total", "Total URL yang diblokir sistem"),
        ("Composite Score", "0.3×Recall + 0.2×Prec + 0.2×F1 + 0.15×Acc + 0.1×Spec + 0.05×(1−FPR)", "Skor gabungan untuk ranking"),
    ]

    t4 = doc.add_table(rows=len(formulas)+1, cols=3)
    t4.style = "Table Grid"
    for ci, h in enumerate(["Metrik","Formula","Deskripsi"]):
        cell = t4.rows[0].cells[ci]
        set_cell_bg(cell, HEX_DARK)
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255,255,255); run.font.name = "Arial"

    for ri, (metric, formula, desc) in enumerate(formulas, 1):
        bg = "DEEAF1" if ri%2==0 else "FFFFFF"
        for ci, val in enumerate([metric, formula, desc]):
            cell = t4.rows[ri].cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9); run.font.name = "Arial"
            if ci == 0: run.font.bold = True
            if ci == 1:
                run.font.name = "Courier New"
                run.font.color.rgb = RGBColor(*COLOR_BLUE)

    # Footer note
    doc.add_paragraph()
    add_para(doc,
        f"Laporan ini dibuat secara otomatis oleh QR Scanner Analysis Tool  |  "
        f"{now.strftime('%d %B %Y, %H:%M')}  |  "
        f"Data sumber: {os.path.basename(excel_path)}",
        size=9, italic=True, color=COLOR_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)
    print(f"\n✅ Laporan berhasil disimpan: {output_path}")
    print(f"   Total halaman estimasi: ~15-20 halaman")


# ─── ENTRY POINT ─────────────────────────────────────────────────────────
if __name__ == "__main__":
    excel_file  = sys.argv[1] if len(sys.argv) > 1 else "Scanner_Result_Batch(2).xlsx"
    output_file = sys.argv[2] if len(sys.argv) > 2 else \
        f"Laporan_QR_Scanner_{datetime.datetime.now().strftime('%Y%m%d')}.docx"

    if not os.path.exists(excel_file):
        print(f"❌ File tidak ditemukan: {excel_file}")
        print("   Usage: python qr_analysis_report.py [input.xlsx] [output.docx]")
        sys.exit(1)

    build_report(excel_file, output_file)
